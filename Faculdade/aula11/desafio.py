
from docx import Document
from docx.shared import Inches


documento_aula = Document()
dis = input("Qual a disciplina: ")
p = documento_aula.add_paragraph("Disciplina: ")
p.add_run(dis).bold = True

carga = input("Qual a Carga horária: ")
p = documento_aula.add_paragraph("Com Carga horária de ")
p.add_run(carga).bold = True
p.add_run(" horas.")

tab = documento_aula.add_table(rows=1, cols=2)
tab.style="Colorful Grid Accent 6"
cels = tab.rows[0].cells
cels[0].text = 'Número'
cels[1].text = "Assunto"


for numero in range(1, 5):
   tarefa = input("Qual é a tarefa de prioridade número-{}? ".format(numero))
   dados = tab.add_row().cells
   dados[0].text = str(numero)
   dados[1].text = tarefa

'''Fazer todos os exercicios 
Assistir 70% das aulas 
Fazer a prova das UAs
Tirar notas acima de MM
'''

documento_aula.save('f:\\z.doc\\doc1.0.docx')

