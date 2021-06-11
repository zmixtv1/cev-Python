
from docx import Document
from docx.shared import Inches

doc = Document()

doc.add_heading('Estes são os seus dados', 0)

nome = input("Qual o seu nome? ")
p = doc.add_paragraph("Seu nome é ")
p.add_run(nome).bold = True
p.add_run('.')

doc.add_heading('A seguir outros detalhes:', level=1)

idade = input("Qual a sua idade? ")
doc.add_paragraph('Você tem {} anos de idade'.format(idade), style='List Bullet')

telefone = input("Qual é o seu telefone? ")
doc.add_paragraph('Seu telefone é o {}'.format(telefone), style='List Bullet')

endereco = input("Qual é o seu endereço? ")
doc.add_paragraph('Seu endereço é {}'.format(endereco), style='List Bullet')

foto = input("Digite o caminho (path) completo de um arquivo com sua foto? ")
doc.add_picture(foto, width=Inches(6), height=Inches(4))

# Aqui entrará o código de criação de tabela

doc.save('f:\\z.doc\\meudoc.docx')

